from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
PURPLE     = RGBColor(0x6B, 0x53, 0xE0)
DARK_GREY  = RGBColor(0x1F, 0x1F, 0x2E)
MID_GREY   = RGBColor(0x44, 0x44, 0x66)
LIGHT_GREY = RGBColor(0x8A, 0x8A, 0xAA)
BLACK      = RGBColor(0x10, 0x10, 0x1A)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_border(cell, sides=("top","bottom","left","right"), color="C0C0D8", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def no_space_after(para):
    pPr = para._p.get_or_add_pPr()
    el  = OxmlElement("w:spacing")
    el.set(qn("w:after"), "0")
    pPr.append(el)

def add_heading(text, level=1):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        run.font.size  = Pt(24)
        run.font.bold  = True
        run.font.color.rgb = DARK_GREY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.bold  = True
        run.font.color.rgb = PURPLE
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(4)
        # left border accent via shading not possible in paragraph — use a thin table trick
    elif level == 3:
        run.font.size  = Pt(11)
        run.font.bold  = True
        run.font.color.rgb = MID_GREY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(text, italic=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = color if color else BLACK
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.bold  = True
        rb.font.size  = Pt(10.5)
        rb.font.color.rgb = BLACK
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(3)
    return p

def add_divider():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D0D0E8")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def add_table(headers, rows, header_bg="2D1F6E"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9.5)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        no_space_after(p)

    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = "F7F6FF" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell_border(cell, color="D8D8EC", sz="4")
            p   = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size  = Pt(9.5)
            run.font.color.rgb = BLACK
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            no_space_after(p)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# ═════════════════════════════════════════════════════════════════════════════
# COVER BLOCK
# ═════════════════════════════════════════════════════════════════════════════
# Purple banner via a 1-row table
banner = doc.add_table(rows=1, cols=1)
banner.alignment = WD_TABLE_ALIGNMENT.LEFT
c = banner.cell(0, 0)
set_cell_bg(c, "2D1F6E")
p  = c.paragraphs[0]
r  = p.add_run("Expense on Demand")
r.font.size  = Pt(26)
r.font.bold  = True
r.font.color.rgb = WHITE
p.alignment  = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(2)

p2 = c.add_paragraph()
r2 = p2.add_run("Phase 3 Analytics Vision — Customer Intelligence & Product Health Platform")
r2.font.size  = Pt(12)
r2.font.color.rgb = RGBColor(0xBB, 0xAA, 0xFF)
p2.alignment  = WD_ALIGN_PARAGRAPH.LEFT
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(14)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Meta line
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
for label, val in [("Prepared for: ", "Management Review    "), ("Date: ", "March 2026")]:
    rb = meta.add_run(label); rb.font.bold = True; rb.font.size = Pt(10); rb.font.color.rgb = MID_GREY
    rv = meta.add_run(val);   rv.font.size = Pt(10); rv.font.color.rgb = BLACK
meta.paragraph_format.space_after = Pt(12)

add_divider()

# ═════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
add_heading("Executive Summary", 2)
add_body(
    "Over the past two phases, we have built an internal analytics platform that gives our delivery teams "
    "real-time visibility into sprint health, bug trends, QA output, team capacity, and release readiness. "
    "That foundation is now mature and operational."
)
add_body(
    "Phase 3 is about turning that analytical capability outward — toward our customers.",
    italic=True, color=PURPLE
)
add_body(
    "The goal is to build a Customer Intelligence layer that tells us, at any point in time: which companies "
    "are thriving on Expense on Demand, which are quietly slipping away, where our product is creating friction, "
    "and what actions we can take — proactively — before a customer churns, escalates, or disengages."
)
add_body(
    "This is not a reporting exercise. It is a strategic shift from reactive customer management to "
    "data-driven account intelligence."
)

# ═════════════════════════════════════════════════════════════════════════════
# 2. WHERE WE ARE TODAY
# ═════════════════════════════════════════════════════════════════════════════
add_heading("Where We Are Today", 2)
add_table(
    ["Phase", "Focus", "Status"],
    [
        ["Phase 1", "Internal delivery metrics — sprint tracking, burndown, team velocity", "Complete"],
        ["Phase 2", "Quality & release health — bug trends, QA output, capacity planning", "Complete"],
        ["Phase 3", "Customer intelligence — product usage, churn risk, approval health", "Proposed"],
    ]
)
add_body("Our current platform answers: \"How is our team performing?\"")
add_body("Phase 3 answers: \"How are our customers performing — and what does that mean for us?\"", italic=True, color=PURPLE)

# ═════════════════════════════════════════════════════════════════════════════
# 3. BUSINESS CASE
# ═════════════════════════════════════════════════════════════════════════════
add_heading("The Business Case", 2)
add_body(
    "Expense management is a workflow product. Customers do not evaluate it on features alone — they evaluate "
    "it on whether it makes expense submission and reimbursement fast and effortless for their employees. "
    "If approval cycles are slow, if employees abandon claims halfway through, if finance teams are manually "
    "chasing down rejections — the product feels broken even if it isn't."
)
add_body("We currently have no systematic visibility into any of this at a customer level.")

add_heading("The Risk of Flying Blind", 3)
add_bullet(
    " A customer's submission volume drops 60% over three months. We find out when they submit a cancellation request.",
    bold_prefix="Scenario 1:"
)
add_bullet(
    " A mid-market client has a 14-day average approval cycle because two managers are not actioning claims. "
    "Their employees complain to HR. HR blames the software.",
    bold_prefix="Scenario 2:"
)
add_bullet(
    " A new client goes live but only 15% of their employees ever submit a claim. Low adoption leads to low "
    "perceived value. They do not renew.",
    bold_prefix="Scenario 3:"
)
add_body(
    "Each of these scenarios is detectable weeks or months in advance with the right data. Phase 3 builds that detection.",
    italic=True
)

# ═════════════════════════════════════════════════════════════════════════════
# 4. VISION
# ═════════════════════════════════════════════════════════════════════════════
add_heading("Vision: What We Are Building", 2)
add_body(
    "A Customer Health Dashboard integrated into our existing analytics platform, providing five core capabilities:"
)

sections = [
    (
        "1. Company Health Scorecards",
        "A per-client view showing submission volume trends, user adoption rates, approval cycle times, and "
        "reimbursement SLA compliance. Each account gets a health score — green, amber, or red — updated on a "
        "rolling basis.",
        "The output: Customer Success and Account Management walk into every QBR knowing exactly which metrics "
        "are healthy and which need a conversation."
    ),
    (
        "2. Churn Risk Detection",
        "Automated flagging of accounts showing early churn signals — declining submission volume, falling active "
        "user counts, or unusually high rejection rates. Flagged accounts surface on a watchlist with the specific "
        "signal that triggered them.",
        "The output: CS teams intervene before the customer has even considered leaving."
    ),
    (
        "3. Approval Chain Intelligence",
        "Company-level visibility into where the approval process breaks down. Which companies have the slowest "
        "manager approval cycles? Where are claims sitting for more than 5 days? Where are rejection rates high — and why?",
        "The output: A targeted conversation with the client admin or a product improvement insight."
    ),
    (
        "4. Product Funnel Analysis",
        "Where in the claim submission flow do employees drop off? Do mobile users complete at the same rate as "
        "web users? Which expense categories generate the most rejections?",
        "The output: Direct input to the product roadmap — validated by real usage data, not assumptions."
    ),
    (
        "5. Release Impact Correlation",
        "Connecting our internal release data (which we already track) with customer-facing outcomes — did "
        "support ticket volume spike after a release? Did approval cycle times change?",
        "The output: Closes the loop between engineering decisions and customer experience."
    ),
]

for title, desc, output in sections:
    add_heading(title, 3)
    add_body(desc)
    p   = doc.add_paragraph()
    rb  = p.add_run("Output: ")
    rb.font.bold  = True
    rb.font.size  = Pt(10.5)
    rb.font.color.rgb = PURPLE
    rv  = p.add_run(output)
    rv.font.size  = Pt(10.5)
    rv.font.italic = True
    rv.font.color.rgb = MID_GREY
    p.paragraph_format.space_after = Pt(6)

# ═════════════════════════════════════════════════════════════════════════════
# 5. DATA REQUIREMENTS
# ═════════════════════════════════════════════════════════════════════════════
add_heading("Data Requirements", 2)
add_body(
    "The following data is required from the relevant engineering and data teams. The majority exists in the "
    "Expense on Demand transactional database today — no new data collection is needed to begin Phase 3a."
)

add_heading("Priority 1 — Core (required to start)", 3)
add_table(
    ["Data Point", "Source", "Notes"],
    [
        ["Claim records with timestamps",    "EoD App DB", "Submitted at, approved at, paid at, rejected at"],
        ["Claim status history",             "EoD App DB", "Full status trail per claim, not just current state"],
        ["Rejection reason codes",           "EoD App DB", "Structured reason codes, not free text"],
        ["Company / tenant identifiers",     "EoD App DB", "Company ID linked to each user and claim"],
        ["User records per company",         "EoD App DB", "Total registered vs. active (submitted in last 30 days)"],
        ["Approver assignment per claim",    "EoD App DB", "Which manager is responsible for each pending claim"],
    ]
)
add_body(
    "Delivery ask: Read-only access to a reporting replica of the EoD production database, or a scheduled daily/hourly "
    "export of the above tables into our analytics environment.", italic=True
)

add_heading("Priority 2 — Behavioural / Funnel (for drop-off analysis)", 3)
add_table(
    ["Data Point", "Source", "Notes"],
    [
        ["Claim draft / session start events", "App instrumentation", "Event listeners at claim creation entry points"],
        ["Receipt upload attempts & outcomes", "App instrumentation", "Success, failure, reason"],
        ["Mobile vs. web session split",       "App instrumentation", "Per user, per session"],
        ["Feature interaction events",         "App instrumentation", "Which features used, by whom, how often"],
    ]
)
add_body(
    "Delivery ask: Instrumentation of key user journey touchpoints using an agreed event schema. Events routed to "
    "a collector we will provide.", italic=True
)

add_heading("Priority 3 — Support Correlation (enhances churn model)", 3)
add_table(
    ["Data Point", "Source", "Notes"],
    [
        ["Support ticket volume per company", "Zendesk / Freshdesk API", "Open date, category, resolution time"],
        ["NPS or CSAT scores per company",    "Survey tool (if in use)", "Linked by company ID"],
    ]
)
add_body("Delivery ask: API credentials for support platform. We will write the integration.", italic=True)

# ═════════════════════════════════════════════════════════════════════════════
# 6. WHAT WE ARE NOT BUILDING
# ═════════════════════════════════════════════════════════════════════════════
add_heading("What We Are NOT Building", 2)
for item in [
    "A CRM replacement — Salesforce/HubSpot remains the record of commercial relationships.",
    "A billing or revenue analytics tool.",
    "An external-facing customer portal — this is an internal intelligence tool.",
    "A replacement for Customer Success workflows — we are informing them, not replacing them.",
]:
    add_bullet(item)

# ═════════════════════════════════════════════════════════════════════════════
# 7. ROADMAP
# ═════════════════════════════════════════════════════════════════════════════
add_heading("Proposed Roadmap", 2)
add_table(
    ["Phase", "Scope", "Data Needed", "Timeline"],
    [
        ["3a — Company Health & Churn",   "Company scorecards, churn watchlist, approval chain visibility", "Priority 1 only",         "Weeks 1–6"],
        ["3b — Product Funnel Analytics", "Drop-off analysis, mobile/web split, feature adoption heatmaps",  "Priority 1 + 2",          "Weeks 7–14"],
        ["3c — Release Impact",          "Engineering ↔ customer outcome correlation, support integration",   "Priority 1 + 2 + 3",      "Weeks 15–20"],
    ]
)
add_body(
    "Phase 3a can begin as soon as Priority 1 database access is granted. No new infrastructure is required — "
    "we extend the platform we have already built."
)

# ═════════════════════════════════════════════════════════════════════════════
# 8. WHAT WE NEED
# ═════════════════════════════════════════════════════════════════════════════
add_heading("What We Need to Proceed", 2)
for i, item in enumerate([
    ("Management sign-off",           "on Phase 3 scope and roadmap"),
    ("Data / engineering engagement", "to provide Priority 1 database access or export pipeline"),
    ("CS and AM alignment",           "define what a churn risk flag should trigger operationally"),
    ("Instrumentation agreement",     "engineering team to agree on event schema for Phase 3b"),
], 1):
    p  = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    rn = p.add_run(f"{i}.  {item[0]}: "); rn.font.bold = True; rn.font.size = Pt(10.5); rn.font.color.rgb = DARK_GREY
    rv = p.add_run(item[1]); rv.font.size = Pt(10.5); rv.font.color.rgb = BLACK

add_divider()

# ═════════════════════════════════════════════════════════════════════════════
# CLOSING STATEMENT
# ═════════════════════════════════════════════════════════════════════════════
closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = closing.add_run(
    "We have the platform. We have the pattern. We have done this for our own teams across two phases and it works.\n"
    "Phase 3 applies the same rigour to the customers who pay for our product.\n"
    "The data is already there. We just need to connect to it."
)
rc.font.size  = Pt(11)
rc.font.italic = True
rc.font.color.rgb = PURPLE
closing.paragraph_format.space_before = Pt(14)
closing.paragraph_format.space_after  = Pt(14)

add_divider()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = footer.add_run("Expense on Demand  ·  Release Analytics Team  ·  Internal Document  ·  March 2026")
rf.font.size  = Pt(8.5)
rf.font.color.rgb = LIGHT_GREY

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(r"c:\Python\Release\Phase3_CustomerAnalytics_Vision.docx")
print("Done.")
